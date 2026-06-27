"""Export an approved ThemeReport as a shareable document.

Produces a Word ``.docx`` when ``python-docx`` is installed, and otherwise a
self-contained HTML file that Word/Pages/Google Docs open cleanly. Either way
the researcher gets the themes, summaries, and supporting quotes (with speaker,
source file, and timecode) as a document they can hand off.
"""

from __future__ import annotations

import html
import io

from ..models import ThemeReport


def _timecode(ms: int) -> str:
    """Compact mm:ss (or h:mm:ss) for human-readable references."""
    total = int(ms) // 1000
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


def _source_name(path: str) -> str:
    return path.replace("\\", "/").rsplit("/", 1)[-1]


def build_docx(report: ThemeReport) -> bytes:
    """Render the report as a .docx. Raises ImportError if python-docx is absent."""
    from docx import Document  # local import so the package stays optional
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(report.display_title, level=0)
    intro = doc.add_paragraph("Focus group themes and supporting quotes")
    intro.runs[0].italic = True

    def add_quote(quote, level: int) -> None:
        bullet = doc.add_paragraph(style="List Bullet 2" if level else "List Bullet")
        run = bullet.add_run(f"“{quote.quote}”")
        run.italic = True
        meta = doc.add_paragraph()
        meta_run = meta.add_run(
            f"{quote.speaker} — {_source_name(quote.source_video)} "
            f"@ {_timecode(quote.start_ms)}–{_timecode(quote.end_ms)}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for i, theme in enumerate(report.themes, start=1):
        doc.add_heading(f"{i}. {theme.title}", level=1)
        if theme.summary:
            doc.add_paragraph(theme.summary)
        for sub in theme.subthemes:
            doc.add_heading(sub.title, level=2)
            for quote in sub.quotes:
                add_quote(quote, level=1)
        for quote in theme.quotes:  # flat quotes (themes without sub-themes)
            add_quote(quote, level=0)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_doc_html(report: ThemeReport) -> str:
    """Dependency-free fallback: an HTML document Word opens as a .doc."""
    esc = html.escape
    parts = [
        "<html><head><meta charset='utf-8'><style>"
        "body{font-family:Calibri,Arial,sans-serif;max-width:760px;margin:40px auto;color:#222}"
        "h1{font-size:24px} h2{font-size:18px;margin-top:28px} h3{font-size:15px;margin:18px 0 4px}"
        ".summary{color:#444} blockquote{margin:8px 0;font-style:italic;"
        "border-left:3px solid #888;padding-left:12px}"
        ".meta{color:#888;font-size:12px;margin:2px 0 12px 12px}"
        ".sub{margin-left:18px}"
        "</style></head><body>",
        f"<h1>{esc(report.display_title)}</h1>",
        "<p class='summary'><em>Focus group themes and supporting quotes</em></p>",
    ]

    def quote_html(quote, sub: bool) -> str:
        cls = "sub" if sub else ""
        return (
            f"<blockquote class='{cls}'>&ldquo;{esc(quote.quote)}&rdquo;</blockquote>"
            f"<div class='meta {cls}'>{esc(quote.speaker)} &middot; "
            f"{esc(_source_name(quote.source_video))} &middot; "
            f"{_timecode(quote.start_ms)}&ndash;{_timecode(quote.end_ms)}</div>"
        )

    for i, theme in enumerate(report.themes, start=1):
        parts.append(f"<h2>{i}. {esc(theme.title)}</h2>")
        if theme.summary:
            parts.append(f"<p class='summary'>{esc(theme.summary)}</p>")
        for sub in theme.subthemes:
            parts.append(f"<h3>{esc(sub.title)}</h3>")
            for quote in sub.quotes:
                parts.append(quote_html(quote, sub=True))
        for quote in theme.quotes:
            parts.append(quote_html(quote, sub=False))
    parts.append("</body></html>")
    return "".join(parts)
